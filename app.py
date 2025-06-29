import streamlit as st
from main import get_cover_letter
from docx import Document
import os

# Custom CSS for improved design
st.markdown(
    """
    <style>
    /* Change text color to white */
    body, h1, h2, h3, h4, h5, h6, p, span {
        color: #FFFFFF !important;
    }

    /* Customize the sidebar */
    [data-testid="stSidebar"] {
        background-color: #2c2c2e;
        color: #FFFFFF;
    }

    /* Customize the input boxes */
    .stTextArea, .stTextInput, .stFileUploader {
        border: 1px solid #FFFFFF !important;
        border-radius: 10px;
        background-color: #333333 !important;
        color: #FAFAFA !important;
        padding: 12px !important;
    }

    /* Increase spacing between text and border in input fields */
    .stTextArea textarea, .stTextInput input, .stFileUploader input {
        padding-top: 10px !important;
        padding-bottom: 10px !important;
    }

    /* Increase font size for labels */
    .stTextArea label, .stTextInput label, .stFileUploader label {
        font-size: 18px !important;
        font-weight: bold !important;
    }

    /* Increase font size for text inside input areas */
    .stTextArea textarea, .stTextInput input, .stFileUploader input {
        font-size: 16px !important;
    }

    /* Customize form submit button */
    .stButton > button {
        background-color: #4CAF50 !important;
        color: white !important;
        font-size: 18px !important;
        padding: 10px 24px !important;
        border-radius: 8px;
        border: none;
    }
    .stButton > button:hover {
        background-color: #45a049 !important;
        color: white !important;
    }

    /* Customize warning message */
    .css-1y0tads {
        background-color: #ffcc00 !important;
        color: #333 !important;
        border: 1px solid #ffa500 !important;
        border-radius: 8px !important;
        padding: 10px !important;
    }

    /* Title styling */
    h1 {
        color: #FFFFFF;
        font-size: 40px;
        font-weight: bold;
    }

    /* Customize form container */
    .stForm {
        background-color: #2c2c2e;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.3);
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title('Cover Letter Document creation using Google Gemini LLM')

gemini_api_key = st.sidebar.text_input('Gemini API Key')

def generate_response(url, cv="",instructions="", gemini_api_key=""):
  print("Generating response",cv)
  output = get_cover_letter(url, cv,instructions, gemini_api_key)
  
  # Create a new Word document  
  doc = Document()
  doc.add_heading('Generated Cover Letter', level=1)
  for paragraph in output.split("\n\n"):
    doc.add_paragraph(paragraph)
    
  # Save the document
  os.makedirs("generated-letters", exist_ok=True)
  save_path = "generated-letters/GeneratedCoverLetter.docx"
  doc.save(save_path)
  st.write(f"Cover letter saved as '{save_path}'")
  st.write(output)

with st.form('my_form'):
    text = st.text_area('Enter LinkedIn Job URL or a job description:', 'Full URL with https:// ...')
    instructions = st.text_area("Enter any specific instructions (eg: short letter, anonymous letter etc.)")
    files = st.file_uploader("Upload one file", type=["pdf"], accept_multiple_files=False)

    submitted = st.form_submit_button('Submit')

    if not gemini_api_key:
        st.warning('Please enter your Gemini API key!', icon='⚠')

    if submitted:
        # Check if the URL starts with "https://"; if not, prepend it
        #if text and not text.startswith("https://"):
        #    text = "https://" + text
        
        generate_response(text, files, instructions, gemini_api_key)
